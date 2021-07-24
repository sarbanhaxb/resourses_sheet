import sys
from PyQt5.QtWidgets import QMainWindow, QApplication, QFileDialog, QMessageBox
import docx
from UI import Ui_main_window
import pandas as pd
import openpyxl
from docx.shared import Inches, Pt

class resourses_sheet(QMainWindow):
    def __init__(self):
        super(resourses_sheet, self).__init__()
        self.UI = Ui_main_window()
        self.UI.setupUi(self)
        self.init_UI()

    def make_magic(self, fname):
        book = pd.read_excel(fname, engine='openpyxl')
        column_one = book.columns.tolist()[0]
        column_two = book.columns.tolist()[1]
        column_three = book.columns.tolist()[2]
        column_four = book.columns.tolist()[3]
        num_slice = []
        for i in range(book.shape[0]):
            if str(book[column_one][i]).replace(" ", '').lower() == 'материалы' or str(book[column_two][i]).replace(" ", '').lower() == 'материалы' or str(
                    book[column_three][i]).replace(" ", '').lower() == 'материалы':
                num_slice.append(i)
            elif str(book[column_one][i]).replace(" ", '').lower() == 'итого"материалы"' or str(
                    book[column_two][i]).replace(" ", '').lower() == 'итого"материалы"':
                num_slice.append(i)
                break
        ws = []
        for i in range(num_slice[0] + 1, num_slice[1]):
            ws.append([str(book[column_one][i]), str(book[column_two][i]), str(book[column_three][i]),
                       str(book[column_four][i])])
        df = pd.DataFrame(ws, columns=['code', 'name', 'unit', 'amount'])

        electrodes = 0 #электроды
        paints = 0 #ЛКМ
        geosynthetic_nt = 0 #геосинтетическая мембрана нетканная
        geogrids_nt = 0 #георешетка нетканная
        thermal = 0 #теплоизоляционный материал
        bituminous = 0 #битумно-резиновая мастика
        sands = 0 #песок
        geomembranes = 0 #геомембрана
        slabs = 0 #дорожные плиты
        semens = 0 #семена
        fertilized = 0 #удобрения
        wire = 0 #провода
        steel_tubes = 0 #стальные трубы

        for index, row in df.iterrows():
            """расчет электродов"""
            if str(row['name']).lower().find('электрод') == 0:
                if str(row['unit']).lower() == 'кг':
                    electrodes += float(row['amount']) / 1000
                elif str(row['unit']).lower() == 'т':
                    electrodes += float(row['amount'])
            """расчет геополотна нетканого"""
            if (str(row['code']).lower().find('фссц-01.7.12.05') == 0 or str(row['code']).lower().find('01.7.12.05') == 0) and (str(row['name']).lower().find('нетканый') == 0 or str(row['name']).lower().find('геополотно') == 0 or str(row['name']).lower().find('геотекстиль') == 0 or str(row['name']).lower().find('дренажный композит') == 0 or str(row['name']).lower().find('материал геотекстильный нетканый') == 0 or str(row['name']).lower().find('полотно иглопробивное') == 0 or str(row['name']).lower().find('полотно нетканное ') == 0):
                if str(row['unit']).lower() == 'м2':
                    geosynthetic_nt += float(row['amount'])
            """расчет георешетки нетканой """
            if (str(row['code']).lower().find('фссц-01.7.12.07') == 0 or str(row['code']).lower().find('01.7.12.07') == 0) and str(row['name']).lower().find('георешетка ') == 0:
                if str(row['unit']).lower() == 'м2':
                    geogrids_nt += float(row['amount'])
            """расчет геомембраны 01.7.12.04"""
            if (str(row['code']).lower().find('01.7.12.04') == 0 or str(row['code']).lower().find('фссц-01.7.12.04') == 0) and (str(row['name']).lower().find('геомембрана') == 0):
                if str(row['unit']).lower() == 'м3':
                    geomembranes += float(row['amount'])
            """расчет теплоизоляционного материала"""
            if (str(row['code']).lower().find('фссц-12.2.04') == 0 or str(row['code']).lower().find('12.2.04') == 0) and (str(row['name']).lower().find('маты') == 0 or str(row['name']).lower().find('пакеты') == 0):
                if str(row['unit']).lower() == 'м3':
                    thermal += float(row['amount'])
            """расчет битумной мастики"""
            if (str(row['code']).lower().find('01.2.03.03') == 0 or str(row['code']).lower().find('фссц-01.2.03.03') == 0) and (str(row['name']).lower().find('мастика ') == 0 or str(row['name']).lower().find('состав мастичный') == 0):
                if str(row['unit']).lower() == 'кг':
                    bituminous += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    bituminous += float(row['amount']) * 1000
            """расчет ЛКМ"""
            #расчет грунтовки группа 14.4.01
            if (str(row['code']).lower().find('14.4.01.') == 0 or str(row['code']).lower().find('фссц-14.4.01.') == 0) and (str(row['name']).lower().find('грунтовка') == 0 or str(row['name']).lower().find('праймер') == 0 or str(row['name']).lower().find('cостав') == 0 or str(row['name']).lower().find('грунт') == 0 or str(row['name']).lower().find('покрытие') == 0):
                if str(row['unit']).lower() == 'кг':
                    paints += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    paints += float(row['amount']) * 1000
            #расчет ЛКМ группы 14.4.02
            if (str(row['code']).lower().find('14.4.02.') == 0 or str(row['code']).lower().find('фссц-14.4.02.') == 0) and (str(row['name']).lower().find('белила') == 0 or str(row['name']).lower().find('краска') == 0 or str(row['name']).lower().find('покрытие') == 0 or str(row['name']).lower().find('эмаль') == 0 or str(row['name']).lower().find('композиция') == 0 or str(row['name']).lower().find('топкоут') == 0):
                if str(row['unit']).lower() == 'кг':
                    paints += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    paints += float(row['amount']) * 1000
            #расчет лаки группы 14.4.03
            if (str(row['code']).lower().find('14.4.03.') == 0 or str(row['code']).lower().find('фссц-14.4.03.') == 0) and (str(row['name']).lower().find('лак') == 0 or str(row['name']).lower().find('раствор хлорсульфированного') == 0 or str(row['name']).lower().find('нитролак') == 0 or str(row['name']).lower().find('покрытие полиуретановое КТ ') == 0):
                if str(row['unit']).lower() == 'кг':
                    paints += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    paints += float(row['amount']) * 1000
            #расчет эмали группы 14.4.04
            if (str(row['code']).lower().find('14.4.04.') == 0 or str(row['code']).lower().find('фссц-14.4.04.') == 0) and (str(row['name']).lower().find('нитроэмаль') == 0 or str(row['name']).lower().find('эмаль') == 0 or str(row['name']).lower().find('покрытие') == 0 or str(row['name']).lower().find('состав (эмаль)') == 0 or str(row['name']).lower().find('шликер') == 0 or str(row['name']).lower().find('спрей-эмаль') == 0 or str(row['name']).lower().find('эпималь') == 0):
                if str(row['unit']).lower() == 'кг':
                    paints += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    paints += float(row['amount']) * 1000
            #другое Прайс-лист 2.1
            if str(row['code']).lower().find('прайс-лист') == 0 and (str(row['name']).lower().find('эмаль') != -1):
                if str(row['unit']).lower() == 'кг':
                    paints += float(row['amount'])
                elif str(row['unit']).lower() == 'т':
                    paints += float(row['amount']) * 1000
            """расчет песка"""
            if str(row['name']).lower().find('песок') == 0:
                if str(row['unit']).lower() == 'м3':
                    sands += float(row['amount'])
                elif str(row['unit']).lower() == 'кг':
                    sands += float(row['amount'])/1600
                elif str(row['unit']).lower() == 'т':
                    sands += float(row['amount'])/1.6
            """расчет щебня"""
            """расчет плит дорожных"""
            if str(row['code']).lower().find('05.1.08.06-0063') != -1 and str(row['name']).lower().find('плиты дорожные') == 0:
                if str(row['unit']).lower() == 'шт':
                    slabs += float(row['amount'])*4.2
            """расчет металлоконструкций"""
            """расчет стальных труб"""
            # Группа 23.3.03.02. Трубы стальные бесшовные горячедеформированные из углеродистой стали
            if str(row['code']).lower().find('23.3.03.02') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.03.02-0001') != -1:
                    steel_tubes += float(row['amount'])*1.48/1000
                elif str(row['code']).lower().find('23.3.03.02-0002') != -1:
                    steel_tubes += float(row['amount']) * 1.78 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0003') != -1:
                    steel_tubes += float(row['amount']) * 4 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0004') != -1:
                    steel_tubes += float(row['amount']) * 7.38 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0005') != -1:
                    steel_tubes += float(row['amount']) * 10.26 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0006') != -1:
                    steel_tubes += float(row['amount']) * 18.99 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0011') != -1:
                    steel_tubes += float(row['amount']) * 1.39 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0012') != -1:
                    steel_tubes += float(row['amount']) * 1.53 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0013') != -1:
                    steel_tubes += float(row['amount']) * 1.63 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0014') != -1:
                    steel_tubes += float(row['amount']) * 1.86 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0015') != -1:
                    steel_tubes += float(row['amount']) * 2.07 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0016') != -1:
                    steel_tubes += float(row['amount']) * 3.08 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0017') != -1:
                    steel_tubes += float(row['amount']) * 1.82 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0018') != -1:
                    steel_tubes += float(row['amount']) * 2.02 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0019') != -1:
                    steel_tubes += float(row['amount']) * 2.15 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0020') != -1:
                    steel_tubes += float(row['amount']) * 2.46 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0021') != -1:
                    steel_tubes += float(row['amount']) * 2.76 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0022') != -1:
                    steel_tubes += float(row['amount']) * 2.62 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0023') != -1:
                    steel_tubes += float(row['amount']) * 2.91 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0024') != -1:
                    steel_tubes += float(row['amount']) * 3.11 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0025') != -1:
                    steel_tubes += float(row['amount']) * 3.58 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0026') != -1:
                    steel_tubes += float(row['amount']) * 4.04 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0027') != -1:
                    steel_tubes += float(row['amount']) * 4.49 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0028') != -1:
                    steel_tubes += float(row['amount']) * 4.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0029') != -1:
                    steel_tubes += float(row['amount']) * 4 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0030') != -1:
                    steel_tubes += float(row['amount']) * 4.62 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0031') != -1:
                    steel_tubes += float(row['amount']) * 5.23 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0032') != -1:
                    steel_tubes += float(row['amount']) * 5.83 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0033') != -1:
                    steel_tubes += float(row['amount']) * 6.41 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0034') != -1:
                    steel_tubes += float(row['amount']) * 6.99 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0035') != -1:
                    steel_tubes += float(row['amount']) * 7.55 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0036') != -1:
                    steel_tubes += float(row['amount']) * 8.63 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0037') != -1:
                    steel_tubes += float(row['amount']) * 6.78 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0038') != -1:
                    steel_tubes += float(row['amount']) * 4.51 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0039') != -1:
                    steel_tubes += float(row['amount']) * 5.22 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0040') != -1:
                    steel_tubes += float(row['amount']) * 5.92 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0041') != -1:
                    steel_tubes += float(row['amount']) * 6.6 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0042') != -1:
                    steel_tubes += float(row['amount']) * 7.27 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0043') != -1:
                    steel_tubes += float(row['amount']) * 7.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0044') != -1:
                    steel_tubes += float(row['amount']) * 8.58 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0045') != -1:
                    steel_tubes += float(row['amount']) * 9.84 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0046') != -1:
                    steel_tubes += float(row['amount']) * 5.4 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0047') != -1:
                    steel_tubes += float(row['amount']) * 6.26 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0048') != -1:
                    steel_tubes += float(row['amount']) * 7.1 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0049') != -1:
                    steel_tubes += float(row['amount']) * 7.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0050') != -1:
                    steel_tubes += float(row['amount']) * 8.75 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0051') != -1:
                    steel_tubes += float(row['amount']) * 9.56 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0052') != -1:
                    steel_tubes += float(row['amount']) * 10.36 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0053') != -1:
                    steel_tubes += float(row['amount']) * 11.91 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0054') != -1:
                    steel_tubes += float(row['amount']) * 6.86 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0055') != -1:
                    steel_tubes += float(row['amount']) * 7.79 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0056') != -1:
                    steel_tubes += float(row['amount']) * 8.71 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0057') != -1:
                    steel_tubes += float(row['amount']) * 9.62 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0058') != -1:
                    steel_tubes += float(row['amount']) * 10.51 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0059') != -1:
                    steel_tubes += float(row['amount']) * 11.39 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0060') != -1:
                    steel_tubes += float(row['amount']) * 13.12 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0061') != -1:
                    steel_tubes += float(row['amount']) * 7.38 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0062') != -1:
                    steel_tubes += float(row['amount']) * 8.38 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0063') != -1:
                    steel_tubes += float(row['amount']) * 9.38 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0064') != -1:
                    steel_tubes += float(row['amount']) * 10.36 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0065') != -1:
                    steel_tubes += float(row['amount']) * 11.33 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0066') != -1:
                    steel_tubes += float(row['amount']) * 12.28 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0067') != -1:
                    steel_tubes += float(row['amount']) * 14.16 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0068') != -1:
                    steel_tubes += float(row['amount']) * 8.5 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0069') != -1:
                    steel_tubes += float(row['amount']) * 9.67 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0070') != -1:
                    steel_tubes += float(row['amount']) * 10.82 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0071') != -1:
                    steel_tubes += float(row['amount']) * 11.96 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0072') != -1:
                    steel_tubes += float(row['amount']) * 13.09 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0073') != -1:
                    steel_tubes += float(row['amount']) * 14.2 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0074') != -1:
                    steel_tubes += float(row['amount']) * 16.4 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0075') != -1:
                    steel_tubes += float(row['amount']) * 10.26 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0076') != -1:
                    steel_tubes += float(row['amount']) * 11.49 / 1000
                elif str(row['code']).lower().find('23.3.03.02-00777') != -1:
                    steel_tubes += float(row['amount']) * 12.7 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0078') != -1:
                    steel_tubes += float(row['amount']) * 13.9 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0079') != -1:
                    steel_tubes += float(row['amount']) * 15.09 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0080') != -1:
                    steel_tubes += float(row['amount']) * 17.44 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0081') != -1:
                    steel_tubes += float(row['amount']) * 19.73 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0082') != -1:
                    steel_tubes += float(row['amount']) * 21.97 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0083') != -1:
                    steel_tubes += float(row['amount']) * 24.17 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0084') != -1:
                    steel_tubes += float(row['amount']) * 10.85 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0085') != -1:
                    steel_tubes += float(row['amount']) * 12.15 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0086') != -1:
                    steel_tubes += float(row['amount']) * 13.44 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0087') != -1:
                    steel_tubes += float(row['amount']) * 14.72 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0088') != -1:
                    steel_tubes += float(row['amount']) * 15.98 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0089') != -1:
                    steel_tubes += float(row['amount']) * 18.47 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0090') != -1:
                    steel_tubes += float(row['amount']) * 20.91 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0091') != -1:
                    steel_tubes += float(row['amount']) * 23.3 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0092') != -1:
                    steel_tubes += float(row['amount']) * 25.65 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0093') != -1:
                    steel_tubes += float(row['amount']) * 11.54 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0094') != -1:
                    steel_tubes += float(row['amount']) * 12.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0095') != -1:
                    steel_tubes += float(row['amount']) * 14.3 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0096') != -1:
                    steel_tubes += float(row['amount']) * 15.67 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0097') != -1:
                    steel_tubes += float(row['amount']) * 17.02 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0098') != -1:
                    steel_tubes += float(row['amount']) * 19.68 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0099') != -1:
                    steel_tubes += float(row['amount']) * 22.29 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0100') != -1:
                    steel_tubes += float(row['amount']) * 24.86 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0101') != -1:
                    steel_tubes += float(row['amount']) * 27.37 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0102') != -1:
                    steel_tubes += float(row['amount']) * 12.13 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0103') != -1:
                    steel_tubes += float(row['amount']) * 13.59 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0104') != -1:
                    steel_tubes += float(row['amount']) * 15.04 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0105') != -1:
                    steel_tubes += float(row['amount']) * 16.48 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0106') != -1:
                    steel_tubes += float(row['amount']) * 17.9 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0107') != -1:
                    steel_tubes += float(row['amount']) * 20.72 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0108') != -1:
                    steel_tubes += float(row['amount']) * 23.48 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0109') != -1:
                    steel_tubes += float(row['amount']) * 26.19 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0110') != -1:
                    steel_tubes += float(row['amount']) * 28.85 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0111') != -1:
                    steel_tubes += float(row['amount']) * 12.72 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0112') != -1:
                    steel_tubes += float(row['amount']) * 14.26 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0113') != -1:
                    steel_tubes += float(row['amount']) * 15.78 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0114') != -1:
                    steel_tubes += float(row['amount']) * 17.29 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0115') != -1:
                    steel_tubes += float(row['amount']) * 18.79 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0116') != -1:
                    steel_tubes += float(row['amount']) * 21.75 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0117') != -1:
                    steel_tubes += float(row['amount']) * 24.66 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0118') != -1:
                    steel_tubes += float(row['amount']) * 27.52 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0119') != -1:
                    steel_tubes += float(row['amount']) * 30.33 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0120') != -1:
                    steel_tubes += float(row['amount']) * 15.7 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0121') != -1:
                    steel_tubes += float(row['amount']) * 17.39 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0122') != -1:
                    steel_tubes += float(row['amount']) * 19.06 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0123') != -1:
                    steel_tubes += float(row['amount']) * 20.72 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0124') != -1:
                    steel_tubes += float(row['amount']) * 23.99 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0125') != -1:
                    steel_tubes += float(row['amount']) * 27.23 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0126') != -1:
                    steel_tubes += float(row['amount']) * 30.41 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0127') != -1:
                    steel_tubes += float(row['amount']) * 33.54 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0128') != -1:
                    steel_tubes += float(row['amount']) * 16.37 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0129') != -1:
                    steel_tubes += float(row['amount']) * 18.13 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0130') != -1:
                    steel_tubes += float(row['amount']) * 19.87 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0131') != -1:
                    steel_tubes += float(row['amount']) * 21.6 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0132') != -1:
                    steel_tubes += float(row['amount']) * 25.03 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0133') != -1:
                    steel_tubes += float(row['amount']) * 28.41 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0134') != -1:
                    steel_tubes += float(row['amount']) * 31.74 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0135') != -1:
                    steel_tubes += float(row['amount']) * 35.02 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0136') != -1:
                    steel_tubes += float(row['amount']) * 17.15 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0137') != -1:
                    steel_tubes += float(row['amount']) * 18.99 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0138') != -1:
                    steel_tubes += float(row['amount']) * 20.82 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0139') != -1:
                    steel_tubes += float(row['amount']) * 22.64 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0140') != -1:
                    steel_tubes += float(row['amount']) * 26.24 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0141') != -1:
                    steel_tubes += float(row['amount']) * 29.79 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0142') != -1:
                    steel_tubes += float(row['amount']) * 33.29 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0143') != -1:
                    steel_tubes += float(row['amount']) * 36.74 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0144') != -1:
                    steel_tubes += float(row['amount']) * 31.57 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0145') != -1:
                    steel_tubes += float(row['amount']) * 22.04 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0146') != -1:
                    steel_tubes += float(row['amount']) * 23.97 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0147') != -1:
                    steel_tubes += float(row['amount']) * 27.79 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0148') != -1:
                    steel_tubes += float(row['amount']) * 35.29 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0149') != -1:
                    steel_tubes += float(row['amount']) * 38.96 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0150') != -1:
                    steel_tubes += float(row['amount']) * 33.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0151') != -1:
                    steel_tubes += float(row['amount']) * 21.58 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0152') != -1:
                    steel_tubes += float(row['amount']) * 23.67 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0153') != -1:
                    steel_tubes += float(row['amount']) * 25.75 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0154') != -1:
                    steel_tubes += float(row['amount']) * 29.86 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0155') != -1:
                    steel_tubes += float(row['amount']) * 37.95 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0156') != -1:
                    steel_tubes += float(row['amount']) * 41.92 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0157') != -1:
                    steel_tubes += float(row['amount']) * 31.52 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0158') != -1:
                    steel_tubes += float(row['amount']) * 36.6 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0159') != -1:
                    steel_tubes += float(row['amount']) * 41.63 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0160') != -1:
                    steel_tubes += float(row['amount']) * 46.61 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0161') != -1:
                    steel_tubes += float(row['amount']) * 51.54 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0162') != -1:
                    steel_tubes += float(row['amount']) * 45.92 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0163') != -1:
                    steel_tubes += float(row['amount']) * 52.28 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0164') != -1:
                    steel_tubes += float(row['amount']) * 58.59 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0165') != -1:
                    steel_tubes += float(row['amount']) * 64.86 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0166') != -1:
                    steel_tubes += float(row['amount']) * 71.07 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0167') != -1:
                    steel_tubes += float(row['amount']) * 77.24 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0168') != -1:
                    steel_tubes += float(row['amount']) * 57.41 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0169') != -1:
                    steel_tubes += float(row['amount']) * 64.36 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0170') != -1:
                    steel_tubes += float(row['amount']) * 71.27 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0171') != -1:
                    steel_tubes += float(row['amount']) * 78.13 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0172') != -1:
                    steel_tubes += float(row['amount']) * 84.93 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0173') != -1:
                    steel_tubes += float(row['amount']) * 62.54 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0174') != -1:
                    steel_tubes += float(row['amount']) * 70.14 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0175') != -1:
                    steel_tubes += float(row['amount']) * 77.68 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0176') != -1:
                    steel_tubes += float(row['amount']) * 85.18 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0177') != -1:
                    steel_tubes += float(row['amount']) * 92.63 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0178') != -1:
                    steel_tubes += float(row['amount']) * 67.67 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0179') != -1:
                    steel_tubes += float(row['amount']) * 75.91 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0180') != -1:
                    steel_tubes += float(row['amount']) * 84.09 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0181') != -1:
                    steel_tubes += float(row['amount']) * 92.23 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0182') != -1:
                    steel_tubes += float(row['amount']) * 100.32 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0183') != -1:
                    steel_tubes += float(row['amount']) * 81.68 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0184') != -1:
                    steel_tubes += float(row['amount']) * 90.51 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0185') != -1:
                    steel_tubes += float(row['amount']) * 99.28 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0186') != -1:
                    steel_tubes += float(row['amount']) * 108.01 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0187') != -1:
                    steel_tubes += float(row['amount']) * 92.55 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0188') != -1:
                    steel_tubes += float(row['amount']) * 102.59 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0189') != -1:
                    steel_tubes += float(row['amount']) * 112.58 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0190') != -1:
                    steel_tubes += float(row['amount']) * 122.51 / 1000
                elif str(row['code']).lower().find('23.3.03.02-0191') != -1:
                    steel_tubes += float(row['amount']) * 113.44 / 1000
            # Группа 23.3.06.01. Трубы стальные водогазопроводные оцинкованные легкие
            elif str(row['code']).lower().find('23.3.06.01') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.01-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.77 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0002') != -1:
                    steel_tubes += float(row['amount']) * 1.08 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0003') != -1:
                    steel_tubes += float(row['amount']) * 1.53 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0004') != -1:
                    steel_tubes += float(row['amount']) * 2.02 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0005') != -1:
                    steel_tubes += float(row['amount']) * 2.74 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0006') != -1:
                    steel_tubes += float(row['amount']) * 3.48 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0007') != -1:
                    steel_tubes += float(row['amount']) * 4.88 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0008') != -1:
                    steel_tubes += float(row['amount']) * 6.6 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0009') != -1:
                    steel_tubes += float(row['amount']) * 7.47 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0010') != -1:
                    steel_tubes += float(row['amount']) * 9.47 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0011') != -1:
                    steel_tubes += float(row['amount']) * 11.94 / 1000
                elif str(row['code']).lower().find('23.3.06.01-0012') != -1:
                    steel_tubes += float(row['amount']) * 14.4 / 1000
            # Группа 23.3.06.02. Трубы стальные водогазопроводные оцинкованные обыкновенные
            elif str(row['code']).lower().find('23.3.06.02') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.02-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.84 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0002') != -1:
                    steel_tubes += float(row['amount']) * 1.19 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0003') != -1:
                    steel_tubes += float(row['amount']) * 1.72 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0004') != -1:
                    steel_tubes += float(row['amount']) * 2.27 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0005') != -1:
                    steel_tubes += float(row['amount']) * 3.15 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0006') != -1:
                    steel_tubes += float(row['amount']) * 4.01 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0007') != -1:
                    steel_tubes += float(row['amount']) * 6.02 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0008') != -1:
                    steel_tubes += float(row['amount']) * 7.5 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0009') != -1:
                    steel_tubes += float(row['amount']) * 8.48 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0010') != -1:
                    steel_tubes += float(row['amount']) * 10.6 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0011') != -1:
                    steel_tubes += float(row['amount']) * 13.37 / 1000
                elif str(row['code']).lower().find('23.3.06.02-0012') != -1:
                    steel_tubes += float(row['amount']) * 16.15 / 1000
            # Группа 23.3.06.03. Трубы стальные водогазопроводные оцинкованные усиленные
            elif str(row['code']).lower().find('23.3.06.03') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.03-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.85 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0002') != -1:
                    steel_tubes += float(row['amount']) * 1.33 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0003') != -1:
                    steel_tubes += float(row['amount']) * 2.07 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0004') != -1:
                    steel_tubes += float(row['amount']) * 2.76 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0005') != -1:
                    steel_tubes += float(row['amount']) * 3.55 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0006') != -1:
                    steel_tubes += float(row['amount']) * 5.05 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0007') != -1:
                    steel_tubes += float(row['amount']) * 6.71 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0008') != -1:
                    steel_tubes += float(row['amount']) * 8.38 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0009') != -1:
                    steel_tubes += float(row['amount']) * 9.49 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0010') != -1:
                    steel_tubes += float(row['amount']) * 11.71 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0011') != -1:
                    steel_tubes += float(row['amount']) * 16.21 / 1000
                elif str(row['code']).lower().find('23.3.06.03-0012') != -1:
                    steel_tubes += float(row['amount']) * 19.6 / 1000
            # Группа 23.3.06.04. Трубы стальные водогазопроводные черные легкие
            elif str(row['code']).lower().find('23.3.06.04') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.04-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.19 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0002') != -1:
                    steel_tubes += float(row['amount']) * 0.3 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0003') != -1:
                    steel_tubes += float(row['amount']) * 0.39 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0004') != -1:
                    steel_tubes += float(row['amount']) * 0.77 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0005') != -1:
                    steel_tubes += float(row['amount']) * 0.73 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0006') != -1:
                    steel_tubes += float(row['amount']) * 1.08 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0007') != -1:
                    steel_tubes += float(row['amount']) * 1.02 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0008') != -1:
                    steel_tubes += float(row['amount']) * 1.53 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0009') != -1:
                    steel_tubes += float(row['amount']) * 2.02 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0010') != -1:
                    steel_tubes += float(row['amount']) * 2.74 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0011') != -1:
                    steel_tubes += float(row['amount']) * 3.48 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0012') != -1:
                    steel_tubes += float(row['amount']) * 4.88 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0013') != -1:
                    steel_tubes += float(row['amount']) * 6.6 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0014') != -1:
                    steel_tubes += float(row['amount']) * 7.47 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0015') != -1:
                    steel_tubes += float(row['amount']) * 9.47 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0033') != -1:
                    steel_tubes += float(row['amount']) * 11.94 / 1000
                elif str(row['code']).lower().find('23.3.06.04-0034') != -1:
                    steel_tubes += float(row['amount']) * 14.4 / 1000
            # Группа 23.3.06.05. Трубы стальные водогазопроводные черные обыкновенные
            elif str(row['code']).lower().find('23.3.06.05') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.05-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.84 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0003') != -1:
                    steel_tubes += float(row['amount']) * 1.72 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0004') != -1:
                    steel_tubes += float(row['amount']) * 2.27 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0005') != -1:
                    steel_tubes += float(row['amount']) * 3.15 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0006') != -1:
                    steel_tubes += float(row['amount']) * 4.01 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0007') != -1:
                    steel_tubes += float(row['amount']) * 6.02 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0010') != -1:
                    steel_tubes += float(row['amount']) * 10.6 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0021') != -1:
                    steel_tubes += float(row['amount']) * 1.19 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0022') != -1:
                    steel_tubes += float(row['amount']) * 7.5 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0023') != -1:
                    steel_tubes += float(row['amount']) * 8.48 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0024') != -1:
                    steel_tubes += float(row['amount']) * 13.37 / 1000
                elif str(row['code']).lower().find('23.3.06.05-0025') != -1:
                    steel_tubes += float(row['amount']) * 16.15 / 1000
            # Группа 23.3.06.06. Трубы стальные водогазопроводные черные усиленные
            elif str(row['code']).lower().find('23.3.06.06') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['code']).lower().find('23.3.06.06-0001') != -1:
                    steel_tubes += float(row['amount']) * 0.93 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0002') != -1:
                    steel_tubes += float(row['amount']) * 1.33 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0003') != -1:
                    steel_tubes += float(row['amount']) * 2.07 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0004') != -1:
                    steel_tubes += float(row['amount']) * 2.76 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0005') != -1:
                    steel_tubes += float(row['amount']) * 3.55 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0006') != -1:
                    steel_tubes += float(row['amount']) * 5.05 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0007') != -1:
                    steel_tubes += float(row['amount']) * 6.71 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0008') != -1:
                    steel_tubes += float(row['amount']) * 8.38 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0009') != -1:
                    steel_tubes += float(row['amount']) * 9.49 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0010') != -1:
                    steel_tubes += float(row['amount']) * 11.71 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0011') != -1:
                    steel_tubes += float(row['amount']) * 16.21 / 1000
                elif str(row['code']).lower().find('23.3.06.06-0012') != -1:
                    steel_tubes += float(row['amount']) * 19.6 / 1000
            # Группа 23.5.02.02. Трубы стальные электросварные прямошовные
            elif str(row['code']).lower().find('23.5.02.02') != -1 and str(row['name']).lower().find('трубы') == 0:
                if str(row['unit']).lower().find('т') == 0:
                    steel_tubes += float(row['amount'])
                elif str(row['unit']).lower().find('м') == 0:
                    if str(row['code']).lower().find('23.5.02.02-0001') != -1:
                        steel_tubes += float(row['amount']) * 0.794 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0002') != -1:
                        steel_tubes += float(row['amount']) * 1.489 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0003') != -1:
                        steel_tubes += float(row['amount']) * 2.286 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0004') != -1:
                        steel_tubes += float(row['amount']) * 4.02 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0005') != -1:
                        steel_tubes += float(row['amount']) * 7.427 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0006') != -1:
                        steel_tubes += float(row['amount']) * 10.324 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0007') != -1:
                        steel_tubes += float(row['amount']) * 12.806 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0008') != -1:
                        steel_tubes += float(row['amount']) * 17.255 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0009') != -1:
                        steel_tubes += float(row['amount']) * 19.11 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0021') != -1:
                        steel_tubes += float(row['amount']) * 0.472 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0022') != -1:
                        steel_tubes += float(row['amount']) * 0.689 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0023') != -1:
                        steel_tubes += float(row['amount']) * 0.893 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0025') != -1:
                        steel_tubes += float(row['amount']) * 1.627 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0026') != -1:
                        steel_tubes += float(row['amount']) * 2.159 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0027') != -1:
                        steel_tubes += float(row['amount']) * 2.159 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0028') != -1:
                        steel_tubes += float(row['amount']) * 2.327 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0029') != -1:
                        steel_tubes += float(row['amount']) * 2.755 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0030') != -1:
                        steel_tubes += float(row['amount']) * 2.823 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0031') != -1:
                        steel_tubes += float(row['amount']) * 3.35 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0032') != -1:
                        steel_tubes += float(row['amount']) * 3.381 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0033') != -1:
                        steel_tubes += float(row['amount']) * 4.02 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0034') != -1:
                        steel_tubes += float(row['amount']) * 4.647 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0035') != -1:
                        steel_tubes += float(row['amount']) * 5.261 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0036') != -1:
                        steel_tubes += float(row['amount']) * 5.435 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0037') != -1:
                        steel_tubes += float(row['amount']) * 6.298 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0038') != -1:
                        steel_tubes += float(row['amount']) * 7.148 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0039') != -1:
                        steel_tubes += float(row['amount']) * 6.274 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0040') != -1:
                        steel_tubes += float(row['amount']) * 5.552 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0041') != -1:
                        steel_tubes += float(row['amount']) * 5.934 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0042') != -1:
                        steel_tubes += float(row['amount']) * 6.879 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0043') != -1:
                        steel_tubes += float(row['amount']) * 7.813 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0044') != -1:
                        steel_tubes += float(row['amount']) * 8.734 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0045') != -1:
                        steel_tubes += float(row['amount']) * 9.642 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0046') != -1:
                        steel_tubes += float(row['amount']) * 5.967 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0047') != -1:
                        steel_tubes += float(row['amount']) * 5.967 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0048') != -1:
                        steel_tubes += float(row['amount']) * 6.379 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0049') != -1:
                        steel_tubes += float(row['amount']) * 7.399 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0050') != -1:
                        steel_tubes += float(row['amount']) * 9.401 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0051') != -1:
                        steel_tubes += float(row['amount']) * 9.401 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0052') != -1:
                        steel_tubes += float(row['amount']) * 10.384 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0053') != -1:
                        steel_tubes += float(row['amount']) * 7.788 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0054') != -1:
                        steel_tubes += float(row['amount']) * 7.788 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0055') != -1:
                        steel_tubes += float(row['amount']) * 10.285 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0056') != -1:
                        steel_tubes += float(row['amount']) * 10.285 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0057') != -1:
                        steel_tubes += float(row['amount']) * 12.733 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0058') != -1:
                        steel_tubes += float(row['amount']) * 12.733 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0059') != -1:
                        steel_tubes += float(row['amount']) * 8.233 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0060') != -1:
                        steel_tubes += float(row['amount']) * 9.562 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0061') != -1:
                        steel_tubes += float(row['amount']) * 10.878 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0062') != -1:
                        steel_tubes += float(row['amount']) * 12.183 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0063') != -1:
                        steel_tubes += float(row['amount']) * 13.474 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0064') != -1:
                        steel_tubes += float(row['amount']) * 13.474 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0065') != -1:
                        steel_tubes += float(row['amount']) * 12.757 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0066') != -1:
                        steel_tubes += float(row['amount']) * 15.823 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0067') != -1:
                        steel_tubes += float(row['amount']) * 14.636 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0068') != -1:
                        steel_tubes += float(row['amount']) * 14.636 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0069') != -1:
                        steel_tubes += float(row['amount']) * 16.41 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0070') != -1:
                        steel_tubes += float(row['amount']) * 12.326 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0071') != -1:
                        steel_tubes += float(row['amount']) * 13.456 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0072') != -1:
                        steel_tubes += float(row['amount']) * 15.329 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0073') != -1:
                        steel_tubes += float(row['amount']) * 15.329 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0074') != -1:
                        steel_tubes += float(row['amount']) * 17.189 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0075') != -1:
                        steel_tubes += float(row['amount']) * 19.037 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0076') != -1:
                        steel_tubes += float(row['amount']) * 22.696 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0077') != -1:
                        steel_tubes += float(row['amount']) * 29.866 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0078') != -1:
                        steel_tubes += float(row['amount']) * 29.866 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0079') != -1:
                        steel_tubes += float(row['amount']) * 18.19 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0080') != -1:
                        steel_tubes += float(row['amount']) * 20.15 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0081') != -1:
                        steel_tubes += float(row['amount']) * 24.031 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0082') != -1:
                        steel_tubes += float(row['amount']) * 24.031 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0083') != -1:
                        steel_tubes += float(row['amount']) * 31.646 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0084') != -1:
                        steel_tubes += float(row['amount']) * 21.262 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0085') != -1:
                        steel_tubes += float(row['amount']) * 21.262 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0086') != -1:
                        steel_tubes += float(row['amount']) * 23.864 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0087') != -1:
                        steel_tubes += float(row['amount']) * 26.454 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0088') != -1:
                        steel_tubes += float(row['amount']) * 31.597 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0089') != -1:
                        steel_tubes += float(row['amount']) * 41.733 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0090') != -1:
                        steel_tubes += float(row['amount']) * 41.733 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0091') != -1:
                        steel_tubes += float(row['amount']) * 26.603 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0092') != -1:
                        steel_tubes += float(row['amount']) * 29.872 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0093') != -1:
                        steel_tubes += float(row['amount']) * 33.13 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0094') != -1:
                        steel_tubes += float(row['amount']) * 39.607 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0095') != -1:
                        steel_tubes += float(row['amount']) * 46.035 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0096') != -1:
                        steel_tubes += float(row['amount']) * 52.414 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0097') != -1:
                        steel_tubes += float(row['amount']) * 31.745 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0098') != -1:
                        steel_tubes += float(row['amount']) * 35.658 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0099') != -1:
                        steel_tubes += float(row['amount']) * 39.558 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0100') != -1:
                        steel_tubes += float(row['amount']) * 47.321 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0101') != -1:
                        steel_tubes += float(row['amount']) * 55.035 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0102') != -1:
                        steel_tubes += float(row['amount']) * 62.699 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0103') != -1:
                        steel_tubes += float(row['amount']) * 70.314 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0104') != -1:
                        steel_tubes += float(row['amount']) * 36.888 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0105') != -1:
                        steel_tubes += float(row['amount']) * 41.443 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0106') != -1:
                        steel_tubes += float(row['amount']) * 45.986 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0107') != -1:
                        steel_tubes += float(row['amount']) * 55.035 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0108') != -1:
                        steel_tubes += float(row['amount']) * 64.034 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0109') != -1:
                        steel_tubes += float(row['amount']) * 72.984 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0110') != -1:
                        steel_tubes += float(row['amount']) * 81.885 / 1000
                    elif str(row['code']).lower().find('23.5.02.02-0111') != -1:
                        steel_tubes += float(row['amount']) * 90.736 / 1000



            """расчет цементного раствора"""
            """расчет песчано-гравийного раствора"""
            """расчет семян трав"""
            if str(row['code']).lower().find('16.2.02.07') != -1 and str(row['name']).lower().find('семена') == 0:
                if str(row['unit']).lower() == 'кг':
                    semens += float(row['amount'])*1
                elif str(row['unit']).lower() == 'т':
                    semens += float(row['amount'])/1000
            """расчет удобрений"""
            if str(row['code']).lower().find('16.3.02.') != -1 and (str(row['name']).lower().find('удобрен') == 0 or str(row['name']).lower().find('селитра') == 0):
                if str(row['unit']).lower() == 'кг':
                    fertilized += float(row['amount'])*1
            """расчет проводов"""
            if str(row['code']).lower().find('21.2.01.01') != -1 and str(row['name']).lower().find('провод') == 0:
                """Группа 21.2.01.01. Провода изолированные для воздушных линий электропередач"""
                if str(row['unit']).lower() == '1000 м':
                    if str(row['name']).lower().find('сип-1') != -1:
                        if str(row['name']).lower().find('1x16+1x25-0,6/1') != -1:
                            wire += float(row['amount'])*1000*0.14
                        elif str(row['name']).lower().find('3x16+1x25-0,6/1') != -1:
                            wire += float(row['amount'])*1000*0.271
                        elif str(row['name']).lower().find('3x25+1x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.381
                        elif str(row['name']).lower().find('3x35+1x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.524
                        elif str(row['name']).lower().find('3x50+1x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.727
                        elif str(row['name']).lower().find('3x50+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.727
                        elif str(row['name']).lower().find('3x70+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.923
                        elif str(row['name']).lower().find('3x70+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.994
                        elif str(row['name']).lower().find('3x95+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.195
                        elif str(row['name']).lower().find('3x95+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.266
                        elif str(row['name']).lower().find('3x120+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.466
                        elif str(row['name']).lower().find('3x150+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.741
                        elif str(row['name']).lower().find('3x185+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 2.33
                        elif str(row['name']).lower().find('3x240+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 2.895
                    elif str(row['name']).lower().find('сип-2') != -1:
                        if str(row['name']).lower().find('3х16+1х25-0,6/1') != -1:
                            wire += float(row['amount'])*1000*0.315
                        elif str(row['name']).lower().find('3x16+1x54,6-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.410
                        elif str(row['name']).lower().find('3x16+1x54,6+2x16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.579
                        elif str(row['name']).lower().find('3x16+1x54,6+2x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.812
                        elif str(row['name']).lower().find('3x25+1x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.422
                        elif str(row['name']).lower().find('3х25+54,6-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.526
                        elif str(row['name']).lower().find('3х25+54,6+2х16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.664
                        elif str(row['name']).lower().find('3x25+1x54,6+2x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.896
                        elif str(row['name']).lower().find('3x35+1x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.568
                        elif str(row['name']).lower().find('3x35+1x50+2x16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.708
                        elif str(row['name']).lower().find('3х35+54,6') != -1:
                            wire += float(row['amount']) * 1000 * 0.735
                        elif str(row['name']).lower().find('3x35+1x54,6+2x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.62 #не точно
                        elif str(row['name']).lower().find('3x50+1x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.727
                        elif str(row['name']).lower().find('3x50+1x50+2x16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.866
                        elif str(row['name']).lower().find('3х50+54,6') != -1:
                            wire += float(row['amount']) * 1000 * 0.776
                        elif str(row['name']).lower().find('3x50+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.8
                        elif str(row['name']).lower().find('3x50+1x70+2x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.97
                        elif str(row['name']).lower().find('3х70+54,6-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.99
                        elif str(row['name']).lower().find('3х70+54,6+2х16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.128
                        elif str(row['name']).lower().find('3х70+70') != -1:
                            wire += float(row['amount']) * 1000 * 1.012
                        elif str(row['name']).lower().find('3x70+1x70+2x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.180
                        elif str(row['name']).lower().find('3x70+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.093
                        elif str(row['name']).lower().find('3x95+1x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.170
                        elif str(row['name']).lower().find('3x95+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.266
                        elif str(row['name']).lower().find('3x120+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.549
                        elif str(row['name']).lower().find('3x150+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.799
                        elif str(row['name']).lower().find('3x185+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 2.146
                        elif str(row['name']).lower().find('3x240+1x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 2.65
                        elif str(row['name']).lower().find('4x16+1x25-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.386
                        elif str(row['name']).lower().find('4x25+1x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.525
                    elif str(row['name']).lower().find('сип-3') != -1:
                        if str(row['name']).lower().find('1x35-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.165
                        elif str(row['name']).lower().find('1x35-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.209
                        elif str(row['name']).lower().find('1х50-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.215
                        elif str(row['name']).lower().find('1x50-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.263
                        elif str(row['name']).lower().find('1х70-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.282
                        elif str(row['name']).lower().find('1x70-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.445
                        elif str(row['name']).lower().find('1х95-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.364
                        elif str(row['name']).lower().find('1x95-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.421
                        elif str(row['name']).lower().find('1х120-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.445
                        elif str(row['name']).lower().find('1x120-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.518
                        elif str(row['name']).lower().find('1х150-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.54
                        elif str(row['name']).lower().find('1x150-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.618
                        elif str(row['name']).lower().find('1x185-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.722
                        elif str(row['name']).lower().find('1x185-35') != -1:
                            wire += float(row['amount']) * 1000 * 0.808
                        elif str(row['name']).lower().find('1x240-20') != -1:
                            wire += float(row['amount']) * 1000 * 0.95
                        elif str(row['name']).lower().find('1x240-35') != -1:
                            wire += float(row['amount']) * 1000 * 1.045
                    elif str(row['name']).lower().find('сип-4') != -1:
                        if str(row['name']).lower().find('2x10-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.092
                        elif str(row['name']).lower().find('2х16') != -1:
                            wire += float(row['amount']) * 1000 * 0.136
                        elif str(row['name']).lower().find('2х25-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.194
                        elif str(row['name']).lower().find('2x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.355
                        elif str(row['name']).lower().find('4х16-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.261
                        elif str(row['name']).lower().find('4х25-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.388
                        elif str(row['name']).lower().find('4x35-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.48
                        elif str(row['name']).lower().find('4x50-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.709
                        elif str(row['name']).lower().find('4x70-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 0.981
                        elif str(row['name']).lower().find('4x95-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.295
                        elif str(row['name']).lower().find('4x120-0,6/1') != -1:
                            wire += float(row['amount']) * 1000 * 1.6
            elif str(row['code']).lower().find('21.2.01.02') != -1 and str(row['name']).lower().find('провод') == 0:
                """Группа 21.2.01.02. Провода неизолированные для воздушных линий электропередач"""
                if str(row['unit']).lower() == 'т':
                    wire += float(row['amount'])*1000
            elif str(row['code']).lower().find('21.2.03.05') != -1 and str(row['name']).lower().find('провод') == 0:
                """Группа 21.2.03.05. Провода силовые для электрических установок на напряжение до 450 в"""
                if str(row['unit']).lower() == '1000 м':
                    if str(row['name']).lower().find('апв') != -1:
                        if str(row['name']).lower().find('2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.014 #неточно
                        elif str(row['name']).lower().find('2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.014
                        elif str(row['name']).lower().find('3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.018
                        elif str(row['name']).lower().find('4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.021
                        elif str(row['name']).lower().find('5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.025
                        elif str(row['name']).lower().find('8-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.04
                        elif str(row['name']).lower().find('10-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.043
                        elif str(row['name']).lower().find('16-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.072
                        elif str(row['name']).lower().find('25-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.105
                        elif str(row['name']).lower().find('35-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.137
                        elif str(row['name']).lower().find('50-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.203
                        elif str(row['name']).lower().find('70-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.269
                        elif str(row['name']).lower().find('95-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.375
                        elif str(row['name']).lower().find('120-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.456
                    elif str(row['name']).lower().find('аппв') != -1:
                        if str(row['name']).lower().find('2x2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.031 #не точно
                        elif str(row['name']).lower().find('2x2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.031
                        elif str(row['name']).lower().find('2x3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.031 #не точно
                        elif str(row['name']).lower().find('2x4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.043
                        elif str(row['name']).lower().find('2x5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.043 #не точно
                        elif str(row['name']).lower().find('2x6-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.057  # не точно
                        elif str(row['name']).lower().find('3x2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.048  # не точно
                        elif str(row['name']).lower().find('3x2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.048  # не точно
                        elif str(row['name']).lower().find('3x3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.048  # не точно
                        elif str(row['name']).lower().find('3x4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.065
                        elif str(row['name']).lower().find('3x5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.074
                        elif str(row['name']).lower().find('3x6-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.085
                    elif str(row['name']).lower().find('пв1') != -1:
                        if str(row['name']).lower().find('0,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0085
                        elif str(row['name']).lower().find('0,75-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0111
                        elif str(row['name']).lower().find('1-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0139
                        elif str(row['name']).lower().find('1,2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.017
                        elif str(row['name']).lower().find('1,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0186
                        elif str(row['name']).lower().find('2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.02 # не точно
                        elif str(row['name']).lower().find('2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.03
                        elif str(row['name']).lower().find('3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.038
                        elif str(row['name']).lower().find('4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.045
                        elif str(row['name']).lower().find('5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.062 # не точно
                        elif str(row['name']).lower().find('6-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0668
                        elif str(row['name']).lower().find('8-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.090
                        elif str(row['name']).lower().find('10 мм2') != -1:
                            wire += float(row['amount']) * 1000 * 0.111
                        elif str(row['name']).lower().find('16 мм2') != -1:
                            wire += float(row['amount']) * 1000 * 0.180
                        elif str(row['name']).lower().find('25-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.283
                        elif str(row['name']).lower().find('35-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.381
                        elif str(row['name']).lower().find('50-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.513
                        elif str(row['name']).lower().find('70-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.728
                        elif str(row['name']).lower().find('95-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.973
                    elif str(row['name']).lower().find('пугв') != -1 and str(row['name']).lower().find('пв3') == -1:
                        if str(row['name']).lower().find('1х0,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.010
                        elif str(row['name']).lower().find('1х0,75-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.012
                        elif str(row['name']).lower().find('1х1-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0154
                        elif str(row['name']).lower().find('1х1,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.022
                        elif str(row['name']).lower().find('1х2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0349
                        elif str(row['name']).lower().find('1х4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0511
                        elif str(row['name']).lower().find('1х10-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.122
                        elif str(row['name']).lower().find('1х25-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.280
                        elif str(row['name']).lower().find('1х35-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.397
                        elif str(row['name']).lower().find('1х50-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.558
                        elif str(row['name']).lower().find('1х70-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.761
                        elif str(row['name']).lower().find('1х95-450') != -1:
                            wire += float(row['amount']) * 1000 * 1.026
                    elif str(row['name']).lower().find('пв3') != -1:
                        if str(row['name']).lower().find(' 2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.01 # не точно
                        elif str(row['name']).lower().find(' 3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.03 # не точно
                        elif str(row['name']).lower().find(' 5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.062
                        elif str(row['name']).lower().find(' 6-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.071
                        elif str(row['name']).lower().find(' 8-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.094
                        elif str(row['name']).lower().find('16-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.182
                    elif str(row['name']).lower().find('ппв') != -1:
                        if str(row['name']).lower().find('2х0,75-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0219
                        elif str(row['name']).lower().find('2х1-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0295
                        elif str(row['name']).lower().find('2х1,2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0343
                        elif str(row['name']).lower().find('2х1,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0398
                        elif str(row['name']).lower().find('2х2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0525
                        elif str(row['name']).lower().find('2х2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.062
                        elif str(row['name']).lower().find('2х3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.076
                        elif str(row['name']).lower().find('2х4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0924
                        elif str(row['name']).lower().find('3х0,75-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0332
                        elif str(row['name']).lower().find('3х1-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0446
                        elif str(row['name']).lower().find('3х1,2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.0511
                        elif str(row['name']).lower().find('3х1,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.060
                        elif str(row['name']).lower().find('3х2-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.079
                        elif str(row['name']).lower().find('3х2,5-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.094
                        elif str(row['name']).lower().find('3х3-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.112
                        elif str(row['name']).lower().find('3х4-450') != -1:
                            wire += float(row['amount']) * 1000 * 0.137
            elif str(row['code']).lower().find('21.2.03.09') != -1 and str(row['name']).lower().find('провод') == 0:
                """Группа 21.2.03.09. Провода, не включенные в группы"""
                if str(row['unit']).lower() == 'т':
                    wire += float(row['amount']) * 1000
                elif str(row['unit']).lower() == '1000 м':
                    if str(row['name']).lower().find('апрн') != -1:
                        if str(row['name']).lower().find('1x6-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.038
                        elif str(row['name']).lower().find('1x10-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.059
                        elif str(row['name']).lower().find('1x35-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.168
                    elif str(row['name']).lower().find('прн') != -1:
                        if str(row['name']).lower().find('2,5 мм2') != -1:
                            wire += float(row['amount']) * 1000 * 0.035
                    elif str(row['name']).lower().find('прпвм') != -1:
                        if str(row['name']).lower().find('2x1,2-380') != -1:
                            wire += float(row['amount']) * 1000 * 0.0417
                    elif str(row['name']).lower().find('пргн') != -1:
                        if str(row['name']).lower().find('1x2,5-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.035
                        elif str(row['name']).lower().find('1x4-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.051
                    elif str(row['name']).lower().find('прто') != -1:
                        if str(row['name']).lower().find('1x1,5-660') != -1:
                            wire += float(row['amount']) * 1000 * 0.025

        global materials
        materials = [['Наименование основных строительных конструкций, изделий и материалов', 'Единица измерения', 'Всего']]
        electrodes = ['Электроды', 'кг', round(electrodes*1000, 2)]
        paints = ['Лакокрасочные материалы', 'кг', round(paints, 2)]
        geosynthetic_nt = ['Нетканое геополотно', 'м2', round(geosynthetic_nt, 2)]
        geogrids_nt = ['Нетканая георешетка', 'м2', round(geogrids_nt, 2)]
        thermal = ['Теплоизоляционный материал', 'м3', round(thermal, 2)]
        bituminous = ['Битумно-резиновая мастика', 'кг', round(bituminous, 2)]
        sands = ['Песок', 'м3', round(round(sands, 2), 2)]
        geomembranes = ['Геомембрана', 'м2', round(geomembranes, 2)]
        slabs = ['Плиты дорожные', 'т', round(slabs, 2)]
        semens = ['Семена трав', "кг", round(semens, 2)]
        fertilized = ['Удобрения', "кг", round(fertilized, 2)]
        wire = ['Провод', 'кг', round(wire, 2)]
        steel_tubes = ['Стальные трубы', 'т', round(steel_tubes, 2)]

        if electrodes[2] != 0:
            materials.append(electrodes)
        if paints[2] != 0:
            materials.append(paints)
        if geosynthetic_nt[2] != 0:
            materials.append(geosynthetic_nt)
        if geogrids_nt[2] != 0:
            materials.append(geogrids_nt)
        if thermal[2] != 0:
            materials.append(thermal)
        if bituminous[2] != 0:
            materials.append(bituminous)
        if sands[2] != 0:
            materials.append(sands)
        if geomembranes[2] != 0:
            materials.append(geomembranes)
        if slabs[2] != 0:
            materials.append(slabs)
        if semens[2] != 0:
            materials.append(semens)
        if fertilized[2] != 0:
            materials.append(fertilized)
        if wire[2] != 0:
            materials.append(wire)
        if steel_tubes[2] != 0:
            materials.append(steel_tubes)

    def showDialog(self):
        try:
            fname = QFileDialog.getOpenFileName(self, 'Open file', '', "Все файлы Excel (*.xlsx)")[0]
            try:
                self.make_magic(fname)
                self.UI.label_5.setText('Ведомость загружена')
            except IndexError:
                QMessageBox.critical(self, "Ошибка ", "Выбран неверный формат файла", QMessageBox.Ok)
                self.UI.label_5.setText('')
        except FileNotFoundError:
            pass

    def Save_file(self):
        name_pd = self.UI.name_pd.toPlainText()
        chifr = self.UI.chifr.toPlainText()
        fn = 'Задание отделу ЭиПБ' #название сохраняемого файла
        try:
            doc = docx.Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Romans'
            font.size = Pt(12)
            a = doc.add_paragraph()
            a.add_run('ЗАДАНИЕ ОТДЕЛУ ЭиПБ').bold = True
            a.alignment = 1
            b = doc.add_paragraph(f'Заказ № {chifr}\nСтадия ПД')
            b.alignment = 2
            doc.add_paragraph('')
            table = doc.add_table(rows=3, cols=2, style='Table Grid')
            table.cell(0, 0).text = 'От отдела\n'
            table.cell(0, 1).text = 'СМ'
            table.cell(1, 0).text = 'Отделу\n'
            table.cell(1, 1).text = 'ЭиПБ'
            table.cell(2, 0).text = 'Наименование объекта\n'
            table.cell(2, 1).text = f'{name_pd}'
            table = doc.add_table(rows=1, cols=1, style='Table Grid')
            table.cell(0, 0).text = 'Ведомость материалов для проведения расчетов и оценки негативного воздействия'
            doc.add_paragraph()
            table = doc.add_table(rows=len(materials), cols=3, style='Table Grid')
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(1.0)
            table.rows[0].cells[0].width = Inches(1.0)
            table.columns[0].width = Inches(1.0)
            table.rows[0].cells[1].width = Inches(1.0)
            x = 0
            y = 0
            for row in table.rows:
                for cell in row.cells:
                    cell.text = str(materials[x][y])
                    y += 1
                x += 1
                y = 0
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(10)
            saves = QFileDialog.getSaveFileName(self,
                                                self.tr("Save file"), fn,
                                                self.tr("Doc files (*.docx *.doc)"))[0]
            if saves:
                try:
                    doc.save(saves)
                except PermissionError:
                    QMessageBox.critical(self, "Ошибка ", "Файл не может быть сохранен, так как открыт в другой программе.", QMessageBox.Ok)
        except NameError:
            QMessageBox.critical(self, "Ошибка ", "Не выбран файл", QMessageBox.Ok)

    def init_UI(self):
        self.setWindowTitle("Анализ ресурсной ведомости")
        self.UI.name_pd.setPlaceholderText('Введите название объекта')
        self.UI.chifr.setPlaceholderText('Введите шифр ПД')
        self.UI.action_5.triggered.connect(self.showDialog)
        self.UI.download_v.clicked.connect(self.Save_file)

app = QApplication([])
application = resourses_sheet()
application.show()

sys.exit(app.exec_())
input('Press ENTER to exit')